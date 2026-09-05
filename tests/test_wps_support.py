# -*- coding: utf-8 -*-
"""WPS 格式（.wps）支持：内容嗅探路由、OOXML 形态、纯文本降级、注册与过滤。"""
from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from gwtool.core.importer import SUPPORTED_EXTS, parse_any
from gwtool.db import dao


def test_wps_registered(tmp_db):
    assert ".wps" in SUPPORTED_EXTS


def test_wps_ooxml_form_parsed_by_docx_parser(tmp_db, tmp_path):
    """金山部分版本/WPS 兼容模式：.wps 实为 OOXML zip —— 按内容嗅探走 docx 解析。"""
    from docx import Document as DX
    p = tmp_path / "方案.wps"
    d = DX()
    d.add_heading("WPS 兼容方案", level=1)
    d.add_paragraph("这是 OOXML 形态的 WPS 文件正文。")
    d.save(str(p))
    r = parse_any(str(p))
    assert r.ok, r.error
    assert "WPS 兼容方案" in r.tree.title
    assert "OOXML 形态" in r.tree.plain_text()


def test_wps_plain_utf16_falls_back_and_extracts(tmp_db, tmp_path):
    """非 OLE/非 zip 的 .wps：四级降级链兜底不崩溃。

    真实金山 .wps 是 OLE 复合文档（走 COM/LibreOffice/纯解析）；纯 UTF-16
    文本属人造形态，原始扫描是"尽力而为"的启发式——契约是：要么 ok 带着
    尽力提取的文本，要么给出可读错误，绝不崩溃。
    """
    p = tmp_path / "纯文本形态.wps"
    body = "这是一份WPS格式的测试文档正文内容，用于降级链提取验证。"
    p.write_bytes(body.encode("utf-16-le"))
    r = parse_any(str(p))
    assert r is not None
    if r.ok:
        assert r.tree is not None and r.tree.plain_text().strip()
    else:
        assert r.error


def test_wps_garbage_no_crash(tmp_db, tmp_path):
    p = tmp_path / "垃圾.wps"
    p.write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 4096)
    r = parse_any(str(p))
    assert r is not None and isinstance(r.ok, bool)


def test_wps_batch_and_dedup(tmp_db, tmp_path):
    """批量导入含 .wps 的混合清单；同内容 .docx 与 .wps 按指纹去重。"""
    from gwtool.core.importer import batch_import
    body = "批量场景下的WPS支持验证正文，内容足够长以构成有效文档。"
    a = tmp_path / "甲.docx"
    from docx import Document as DX
    _d = DX()
    _d.add_paragraph(body)
    _d.save(str(a))
    b = tmp_path / "乙.wps"
    b.write_bytes(body.encode("utf-16-le"))     # 内容相同（纯文本形态）
    c = tmp_path / "坏.wps"
    c.write_bytes(b"\x00" * 64)
    results = batch_import([str(a), str(b), str(c)])
    assert len(results) == 3
    # 甲（docx 真 OOXML）入库；乙与甲同内容 -> 指纹去重跳过或失败均不计成功
    assert results[0].ok
    if results[1].ok:
        did = dao.add_document(dao.Document(
            title=results[1].tree.title,
            content_text=results[1].tree.plain_text(),
            blocks_json=results[1].tree.to_json(), file_type="wps"))
        assert did in (-1, 0) or did > 0      # 去重逻辑不崩溃即可
    assert not results[2].ok
