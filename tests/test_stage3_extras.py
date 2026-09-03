# -*- coding: utf-8 -*-
"""阶段3：体检报告导出、智能分类建议、批量打标签。"""
import pytest

from gwtool.core import classify, report
from gwtool.core.inspector import Finding
from gwtool.db import dao


def _add_doc(title: str, content: str, category_id: int = 0, tags: str = "") -> int:
    return dao.add_document(dao.Document(title=title, content_text=content,
                                         category_id=category_id, tags=tags))


# ============================================================ 体检报告
def _findings():
    return [
        Finding(severity="error", item="发文字号", detail="缺少发文字号"),
        Finding(severity="warn", item="成文日期", detail="建议用阿拉伯数字标注"),
        Finding(severity="info", item="页码", detail="页码格式为外侧四字线"),
    ]


def test_summarize_and_verdict():
    assert report.summarize(_findings()) == {"error": 1, "warn": 1, "info": 1}
    assert "1 项不合规范" in report.verdict(_findings())
    assert "未发现硬性错误" in report.verdict(_findings()[1:])
    assert "未发现格式问题" in report.verdict([])


def test_report_tree_orders_by_severity_and_has_table():
    tree = report.build_report_tree(_findings(), source_name="关于xx的通知.docx")
    assert tree.title == "公文格式体检报告"
    tables = [b for b in tree.blocks if b.type == "table"]
    assert len(tables) == 1, "报告应含一张检查明细表"
    rows = tables[0].rows
    assert rows[0] == ["序号", "结论", "检查项", "说明"]
    # 不合规范排最前，便于先看必须整改的
    assert rows[1][2] == "发文字号" and rows[1][1] == "不合规范"
    assert rows[-1][2] == "页码" and rows[-1][1] == "提示"
    assert [r[0] for r in rows[1:]] == ["1", "2", "3"], "序号应连续"


def test_report_tree_without_findings():
    tree = report.build_report_tree([])
    assert not [b for b in tree.blocks if b.type == "table"]
    assert any("未发现问题" in b.text for b in tree.blocks if b.type == "paragraph")


def test_export_report_produces_readable_docx(tmp_db, tmp_path):
    """端到端：导出的 docx 必须能被 python-docx 打开且含结论与明细。"""
    from docx import Document

    out = tmp_path / "体检报告.docx"
    report.export_report(_findings(), str(out), source_name="测试公文")
    assert out.exists() and out.stat().st_size > 5000

    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "公文格式体检报告" in text
    assert "测试公文" in text
    assert "GB/T 9704" in text
    assert "1 项不合规范" in text
    # 明细以表格承载
    assert doc.tables, "报告缺少明细表格"
    cells = [c.text for row in doc.tables[0].rows for c in row.cells]
    assert "发文字号" in cells and "不合规范" in cells


# ============================================================ 智能分类建议
def test_suggest_returns_empty_when_no_categories(tmp_db):
    assert classify.suggest("关于安全生产的通知") == []


def test_suggest_returns_empty_when_categories_have_no_docs(tmp_db):
    dao.add_category("空分类")
    assert classify.suggest("关于安全生产的通知") == []


def test_suggest_picks_the_matching_category(tmp_db):
    safety = dao.add_category("安全生产")
    finance = dao.add_category("财务预算")
    for i in range(3):
        _add_doc(f"安全生产检查通知{i}",
                 "压实安全生产责任，排查隐患，防范各类事故发生，开展安全生产专项整治。",
                 category_id=safety)
    for i in range(3):
        _add_doc(f"预算安排通知{i}",
                 "编制年度财务预算，规范资金使用，严格财务报销与决算管理。",
                 category_id=finance)

    got = classify.suggest("关于深入开展安全生产隐患排查整治工作的通知", top_n=2)
    assert got, "应给出建议"
    assert got[0][0] == safety, f"首推应是安全生产，实际 {got}"
    assert got[0][2] > 0
    names = [n for _c, n, _s in got]
    assert "安全生产" in names


def test_suggest_scores_are_normalised_and_sorted(tmp_db):
    cat = dao.add_category("党建")
    _add_doc("党建工作要点", "加强党的建设，落实主体责任，推进党风廉政建设。", category_id=cat)
    got = classify.suggest("党风廉政建设与党的建设工作要点", top_n=3)
    scores = [s for _c, _n, s in got]
    assert scores == sorted(scores, reverse=True), "得分应降序"
    assert all(0 < s <= 1 for s in scores), f"得分应归一化到 (0,1]，实际 {scores}"


def test_suggest_ignores_unrelated_text(tmp_db):
    cat = dao.add_category("安全生产")
    _add_doc("安全通知", "压实安全生产责任，排查隐患。", category_id=cat)
    # 完全不相干的文本不该硬凑一个建议
    assert classify.suggest("量子纠缠态的贝尔不等式检验") == []


def test_suggest_for_document_weights_title(tmp_db):
    """被错放的文档：建议应指向它内容真正所属的分类，而不是当前所在分类。"""
    other = dao.add_category("财务预算")
    target = dao.add_category("安全生产")
    _add_doc("预算通知", "编制年度财务预算，规范资金使用。", category_id=other)
    for i in range(3):
        _add_doc(f"安全生产检查{i}",
                 "压实安全生产责任，排查安全隐患，防范各类事故发生。",
                 category_id=target)
    # 这篇内容属于安全生产，却被错放在财务预算下
    did = _add_doc("安全生产专项整治方案", "开展安全隐患排查，压实安全生产责任。",
                   category_id=other)

    got = classify.suggest_for_document(did, top_n=2)
    assert got, "应给出建议"
    assert got[0][0] == target, f"首推应是安全生产({target})，实际 {got}"


def test_suggest_for_missing_document(tmp_db):
    assert classify.suggest_for_document(999999) == []


def test_profiles_reusable_across_calls(tmp_db):
    """批量建议时画像只应算一次，故 suggest 必须接受外部传入的 profiles。"""
    cat = dao.add_category("安全生产")
    _add_doc("安全通知", "压实安全生产责任，排查隐患。", category_id=cat)
    profiles = classify.category_profiles()
    assert cat in profiles
    assert classify.suggest("安全生产责任", profiles=profiles)
    assert classify.suggest("安全生产责任", profiles=profiles) == \
           classify.suggest("安全生产责任", profiles=profiles), "同画像应得同结果"


# ============================================================ 批量打标签
def test_bulk_add_tags(tmp_db):
    # 正文必须各不相同：documents 对 text_hash 有唯一索引，内容重复会被去重
    ids = [_add_doc(f"文档{i}", f"这是第{i}篇的独立正文内容", tags="原有")
           for i in range(3)]
    assert len(set(ids)) == 3, f"文档未按预期各自入库：{ids}"
    n = dao.bulk_update_tags(ids, add=("2026年度", "重点"))
    assert n == 3
    for did in ids:
        tags = dao.get_document(did).tags
        assert "原有" in tags and "2026年度" in tags and "重点" in tags


def test_bulk_add_tags_dedupes_and_is_idempotent(tmp_db):
    did = _add_doc("文档", "内容", tags="重点")
    assert dao.bulk_update_tags([did], add=("重点",)) == 0, "已有标签不应算作变更"
    assert dao.get_document(did).tags == "重点"
    dao.bulk_update_tags([did], add=("重点", "2026"))
    assert dao.get_document(did).tags.count("重点") == 1, "标签重复了"


def test_bulk_remove_tags(tmp_db):
    did = _add_doc("文档", "内容", tags="甲，乙，丙")
    n = dao.bulk_update_tags([did], remove=("乙",))
    assert n == 1
    tags = dao.get_document(did).tags
    assert "乙" not in tags
    assert "甲" in tags and "丙" in tags


def test_bulk_tags_handles_mixed_comma_separators(tmp_db):
    """历史数据中英文逗号混用，必须都能正确解析，否则会切出错误的标签。"""
    did = _add_doc("文档", "内容", tags="甲,乙，丙")
    dao.bulk_update_tags([did], remove=("乙",))
    tags = dao.get_document(did).tags
    assert "乙" not in tags
    assert "甲" in tags and "丙" in tags
    assert "," not in tags, "写回应统一为中文逗号"


def test_bulk_tags_add_and_remove_in_one_call(tmp_db):
    did = _add_doc("文档", "内容", tags="旧标签")
    dao.bulk_update_tags([did], add=("新标签",), remove=("旧标签",))
    tags = dao.get_document(did).tags
    assert "新标签" in tags and "旧标签" not in tags


def test_bulk_tags_noop_cases(tmp_db):
    did = _add_doc("文档", "内容", tags="甲")
    assert dao.bulk_update_tags([], add=("乙",)) == 0
    assert dao.bulk_update_tags([did]) == 0, "既不加也不删应为空操作"
    assert dao.bulk_update_tags([did], add=("  ", "")) == 0, "空标签应被忽略"
    assert dao.get_document(did).tags == "甲"


def test_bulk_tags_skips_missing_documents(tmp_db):
    did = _add_doc("文档", "内容")
    n = dao.bulk_update_tags([did, 999999], add=("标签",))
    assert n == 1, "不存在的文档应跳过而不是报错"
