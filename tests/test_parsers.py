# -*- coding: utf-8 -*-
"""解析器测试：动态生成样本文件（docx/txt/rtf/md/html/pdf/doc兜底）。"""
from pathlib import Path

from gwtool.core.importer import batch_import, parse_any


def _make_docx(path: Path):
    from docx import Document
    doc = Document()
    doc.add_heading("关于测试工作的报告", level=1)
    doc.add_paragraph("第一段：这是正文内容，用于验证解析完整性。")
    doc.add_heading("（一）基本情况", level=2)
    doc.add_paragraph("第二段：详细情况说明。")
    doc.save(str(path))


def _make_pdf(path: Path):
    import pymupdf as fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 100), "PDF解析测试文档", fontname="china-s", fontsize=20)
    page.insert_text((72, 140), "这是第一段正文内容，验证文字提取。", fontname="china-s", fontsize=12)
    page.insert_text((72, 160), "这是第二段正文内容。", fontname="china-s", fontsize=12)
    doc.save(str(path))
    doc.close()


def _make_rtf(path: Path):
    content = r"{\rtf1\ansi{\fonttbl{\f0 SimSun;}}\f0\fs28 关于测试的RTF文档\par 第一段内容。\par 第二段内容。\par}"
    path.write_text(content, encoding="gb18030")


def _make_doc_fallback(path: Path):
    """构造一个极简 OLE 复合文件走兜底扫描路径较为复杂；
    直接验证纯文本兜底函数的行为。"""
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)


def test_parse_docx(tmp_path):
    p = tmp_path / "sample.docx"
    _make_docx(p)
    r = parse_any(str(p))
    assert r.ok
    assert r.tree.title == "关于测试工作的报告"
    texts = [b.text for b in r.tree.blocks]
    assert "第一段：这是正文内容，用于验证解析完整性。" in texts
    levels = [b.level for b in r.tree.blocks if b.type == "heading"]
    assert 1 in levels and 2 in levels


def test_parse_txt_gbk(tmp_path):
    p = tmp_path / "sample.txt"
    p.write_bytes("关于测试的文本文档\n第一段中文内容。\n第二段内容。".encode("gbk"))
    r = parse_any(str(p))
    assert r.ok
    assert "第一段中文内容。" in r.tree.plain_text()


def test_parse_md(tmp_path):
    p = tmp_path / "sample.md"
    p.write_text("# 主标题\n\n## 二级标题\n\n正文段落。\n- 列表项一\n- 列表项二\n",
                 encoding="utf-8")
    r = parse_any(str(p))
    assert r.ok
    assert r.tree.title == "主标题"
    assert any(b.type == "heading" and b.level == 2 for b in r.tree.blocks)
    assert any(b.type == "list_item" for b in r.tree.blocks)


def test_parse_html(tmp_path):
    p = tmp_path / "sample.html"
    p.write_text("<html><head><title>网页标题</title></head><body>"
                 "<h1>第一章</h1><p>段落内容一。</p><p>段落内容二。</p>"
                 "</body></html>", encoding="utf-8")
    r = parse_any(str(p))
    assert r.ok
    assert "段落内容一。" in r.tree.plain_text()


def test_parse_rtf(tmp_path):
    p = tmp_path / "sample.rtf"
    _make_rtf(p)
    r = parse_any(str(p))
    assert r.ok
    text = r.tree.plain_text()
    assert "第一段内容。" in text


def test_parse_pdf(tmp_path):
    p = tmp_path / "sample.pdf"
    _make_pdf(p)
    r = parse_any(str(p))
    assert r.ok, r.error
    assert "这是第一段正文内容，验证文字提取。" in r.tree.plain_text()


def test_parse_scanned_pdf_reports_failure(tmp_path):
    import pymupdf as fitz
    p = tmp_path / "empty.pdf"
    doc = fitz.open()
    doc.new_page()
    doc.save(str(p))
    doc.close()
    r = parse_any(str(p))
    assert not r.ok and "未提取到文字" in r.error


def _import_like_worker(files):
    """模拟 ImportWorker 的入库逻辑（解析 + 去重写入）。"""
    from gwtool.db import dao
    ok = skip = 0
    for p in files:
        r = parse_any(p)
        if not r.ok or r.tree is None:
            skip += 1
            continue
        from pathlib import Path
        doc = dao.Document(
            title=r.tree.title or Path(p).stem,
            content_text=r.tree.plain_text(),
            blocks_json=r.tree.to_json(),
            file_path=p,
            file_type=Path(p).suffix.lower().lstrip("."),
        )
        if dao.add_document(doc) < 0:
            skip += 1
        else:
            ok += 1
    return ok, skip


def test_batch_import_dedup(tmp_db, tmp_path):
    p1 = tmp_path / "a.txt"
    p2 = tmp_path / "b.txt"
    p1.write_text("完全相同的测试内容，用于验证去重。", encoding="utf-8")
    p2.write_text("完全相同的测试内容，用于验证去重。", encoding="utf-8")
    results = batch_import([str(p1), str(p2)])
    assert all(r.ok for r in results)
    ok, skip = _import_like_worker([str(p1), str(p2)])
    assert ok == 1 and skip == 1
    from gwtool.db import dao
    assert dao.count_documents() == 1


def test_batch_import_50_files_speed(tmp_db, tmp_path):
    """验收：50个文件批量导入（解析+入库） < 10 秒。"""
    import time
    files = []
    for i in range(50):
        p = tmp_path / f"f{i}.txt"
        p.write_text(f"第{i}份材料标题\n这是第{i}份材料的正文内容，包含足够的文字。" * 20,
                     encoding="utf-8")
        files.append(str(p))
    t0 = time.time()
    ok, skip = _import_like_worker(files)
    elapsed = time.time() - t0
    assert ok == 50 and skip == 0
    assert elapsed < 10, f"批量导入耗时 {elapsed:.1f}s"
