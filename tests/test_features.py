# -*- coding: utf-8 -*-
"""扩展功能测试：工具箱/骨架/快照/体检/查重/水印/忽略名单/加密备份/批量汇编。"""
import os
from pathlib import Path

import pytest

# ================================================================ F2 工具箱
from gwtool.core import toolbox as tb


def test_amount_to_cn_basic():
    assert tb.amount_to_cn("10050000.30") == "人民币壹仟零伍万元叁角整"
    assert tb.amount_to_cn("10050000.00") == "人民币壹仟零伍万元整"
    assert tb.amount_to_cn("105") == "人民币壹佰零伍元整"
    assert tb.amount_to_cn("10.00") == "人民币壹拾元整"
    assert tb.amount_to_cn("0.05") == "人民币伍分"
    assert tb.amount_to_cn("0.00") == "零元整"
    assert tb.amount_to_cn(1234.56) == "人民币壹仟贰佰叁拾肆元伍角陆分"
    assert tb.amount_to_cn("1,234,567.89") == "人民币壹佰贰拾叁万肆仟伍佰陆拾柒元捌角玖分"


def test_amount_to_cn_zero_simplification():
    # 中间连续零只写一个零
    assert "零零" not in tb.amount_to_cn("10000005.00")


def test_date_and_number_cn():
    assert tb.digits_to_cn_date("2026年8月30日发布") == "二〇二六年八月三十日发布"
    assert tb.digits_to_cn_date("自2026-01-05起") == "自二〇二六年一月五日起"
    assert tb.digits_to_cn_date("2026年10月1日") == "二〇二六年十月一日"
    assert tb.number_to_upper_cn("202612") == "贰零贰陆壹贰"


def test_full_half_width():
    assert tb.full_to_half("ａｂｃ１２３，：") == "abc123,:"
    assert tb.half_to_full("abc123") == "ａｂｃ１２３"
    assert tb.full_to_half("中文保留１２３") == "中文保留123"


def test_s2t_t2s():
    assert tb.s2t("简体") == "簡體"
    assert tb.t2s(tb.s2t("公文汇编助手")) == "公文汇编助手"


# ================================================================ F1 骨架
def test_skeletons_15_kinds():
    from gwtool.core.skeletons import kinds
    assert len(kinds()) == 15
    for k in kinds():
        assert k.strip()


def test_skeleton_qingshi_render():
    from gwtool.core.skeletons import get
    text = get("请示").render(title="××单位关于××事项的请示",
                             org="××单位", recipients="上级机关：",
                             matter="××事项", date="2026年8月30日")
    assert "上级机关：" in text
    assert "妥否，请批示。" in text
    assert "2026年8月30日" in text


def test_skeleton_baogao_no_qingshi():
    """报告骨架不得含请示结束语。"""
    from gwtool.core.skeletons import get
    text = get("报告").render(title="关于××的报告", org="×", matter="×",
                             date="2026年8月30日")
    assert "特此报告。" in text
    assert "请批示" not in text and "请批复" not in text


# ================================================================ F3 快照
def test_snapshot_roundtrip_and_prune(tmp_db):
    from gwtool.db import dao
    did = dao.add_document(dao.Document(title="t", content_text="v1"))
    for i in range(35):
        dao.add_snapshot(did, "t", f"版本{i}", reason="auto")
    snaps = dao.list_snapshots(did)
    assert len(snaps) == 30  # 轮转保留 30 份
    newest = dao.get_snapshot(snaps[0]["id"])
    assert newest["content"] == "版本34"
    old = dao.get_snapshot(snaps[-1]["id"])
    assert old["content"] == "版本5"  # 版本0-4 被清理


# ================================================================ F9 忽略名单
def test_ignore_words(tmp_db):
    from gwtool.core import corrector
    from gwtool.db import dao
    dao.add_error_pair("专有名", "其他词", "测试", 0.9, source="user")
    corrector.invalidate_cache()
    text = "这里包含专有名一词。"
    assert any(c.wrong == "专有名" for c in corrector.check_text(text))
    dao.add_ignore_word("专有名")
    corrector.invalidate_cache()
    assert not any(c.wrong == "专有名" for c in corrector.check_text(text))
    dao.remove_ignore_word("专有名")
    corrector.invalidate_cache()
    assert any(c.wrong == "专有名" for c in corrector.check_text(text))


# ================================================================ F8 机构沿革
def test_org_rename_pairs(tmp_db):
    from gwtool.core import corrector
    corr = corrector.check_text("根据环境保护部的要求执行。")
    hits = [c for c in corr if c.wrong == "环境保护部"]
    assert hits and hits[0].category == "机构沿革"
    assert hits[0].suggestion == "生态环境部"


# ================================================================ F7 体检
def test_inspector_text_numbering_and_date(tmp_db):
    from gwtool.core.inspector import inspect_text
    text = ("关于××的通知\n\n各科室：\n"
            "一、总体要求\n"
            "1.跳级内容\n"
            "三、编号跳跃\n"
            "联系人[2026]通知。详见2026.08.30文件。\n")
    findings = inspect_text(text, kind_hint="通知")
    items = [f.item for f in findings]
    assert "标题编号" in items
    assert "标题层级" in items
    assert "发文字号" in items
    assert "成文日期" in items


def test_inspector_baogao_with_qingshi(tmp_db):
    from gwtool.core.inspector import inspect_text
    text = "关于××的报告\n\n上级：\n情况如下。\n请批示。\n"
    findings = inspect_text(text, kind_hint="报告")
    assert any(f.item == "文种混用" and f.severity == "error" for f in findings)


def test_inspector_docx(tmp_db, tmp_path):
    from gwtool.core.inspector import inspect_docx
    from docx import Document
    from docx.shared import Mm, Pt
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Mm(20)   # 不合规边距
    p = doc.add_paragraph("正文内容测试")
    p.runs[0].font.size = Pt(12)  # 非三号
    p2 = doc.add_paragraph("更多正文内容用于采样检查通过阈值判定。" * 20)
    p2.runs[0].font.size = Pt(12)
    path = tmp_path / "bad.docx"
    doc.save(str(path))
    findings = inspect_docx(str(path))
    assert any("页边距" in f.item for f in findings)


# ================================================================ F11 查重
def test_simhash_similar(tmp_db):
    from gwtool.core.simhash import find_similar, similarity
    a = "乡村振兴要坚持农业农村优先发展，巩固拓展脱贫攻坚成果，扎实推进共同富裕各项工作。"
    b = "乡村振兴要坚持农业农村优先发展，巩固拓展脱贫攻坚成果，扎实推进共同富裕各项任务。"
    c = "今天天气晴朗，适合外出郊游，公园里的花朵都开放了，孩子们在草地上奔跑嬉戏。"
    assert similarity(a, b) > 0.8   # 仅两字之差，3-gram Jaccard 约 0.85
    assert similarity(a, c) < 0.3
    pairs = find_similar({1: a, 2: b, 3: c}, threshold=0.7)
    assert (1, 2) in [(p[0], p[1]) for p in pairs] or (2, 1) in [(p[0], p[1]) for p in pairs]
    assert all({p[0], p[1]} != {1, 3} for p in pairs)


# ================================================================ F12 水印
def test_watermark_pdf(tmp_db, tmp_path, qapp):
    from gwtool.core.watermark import stamp_watermark_pdf
    import pymupdf as fitz
    src = tmp_path / "in.pdf"
    d = fitz.open()
    d.new_page()
    d.save(str(src))
    d.close()
    out = stamp_watermark_pdf(str(src), "征求意见稿", str(tmp_path / "wm.pdf"))
    d = fitz.open(out)
    assert "征求意见稿" in d[0].get_text()
    d.close()


def test_watermark_docx(tmp_db, tmp_path):
    from docx import Document
    from gwtool.core.watermark import add_watermark_docx
    path = tmp_path / "in.docx"
    Document().save(str(path))
    out = add_watermark_docx(str(path), "秘密★1年", str(tmp_path / "wm.docx"))
    from docx import Document as D2
    xml = D2(out).sections[0].header._element.xml
    assert "秘密★1年" in xml and "WaterMarkObject" in xml


# ================================================================ F14 安全
def test_password_lock(tmp_db):
    from gwtool.core import security
    assert not security.has_password()
    security.set_password("abcd1234")
    assert security.has_password()
    assert security.verify_password("abcd1234")
    assert not security.verify_password("wrong")
    security.clear_password()
    assert not security.has_password()


def test_encrypted_backup_roundtrip(tmp_db, tmp_path):
    pytest.importorskip("pyzipper")
    import gwtool.core.backup as bk
    from gwtool import paths

    def fake_dir():
        d = tmp_path / "bk"
        d.mkdir(parents=True, exist_ok=True)
        return d

    bk.paths.backup_dir = fake_dir
    from gwtool.db import dao
    dao.add_document(dao.Document(title="加密备份测试", content_text="机密内容。"))
    z = bk.create_backup(note="加密", password="pass1234")
    assert "加密" in Path(z).name
    # 用错误口令恢复应失败
    from gwtool.db.connection import get_conn
    get_conn().execute("DELETE FROM documents")
    get_conn().commit()
    with pytest.raises(Exception):
        bk.restore_backup(z, password="wrongpass")
    # 正确口令恢复成功
    bk.restore_backup(z, password="pass1234")
    d = dao.get_document(1)
    assert d and "机密内容" in d.content_text


def test_backup_rotation(tmp_db, tmp_path):
    import gwtool.core.backup as bk

    def fake_dir():
        d = tmp_path / "bk2"
        d.mkdir(parents=True, exist_ok=True)
        return d

    bk.paths.backup_dir = fake_dir
    for i in range(23):
        (fake_dir() / f"gwtool_backup_{i:08d}_000.zip").write_text("x")
    removed = bk.rotate_backups(keep_recent=20)
    assert removed == 3
    assert len(list(fake_dir().glob("*.zip"))) == 20


# ================================================================ F15 批量汇编
def test_batch_compile(tmp_db, tmp_path):
    from gwtool.core.batch import batch_compile_each
    from gwtool.core.template import default_template
    from gwtool.db import dao
    from gwtool.core.model import Block, DocTree

    ids = []
    for name in ("文件甲", "文件乙", "文件丙"):
        tree = DocTree(title=f"关于{name}的通知",
                       blocks=[Block(text="正文内容。"), Block(text="第二条内容。")])
        ids.append(dao.add_document(dao.Document(
            title=tree.title, content_text=tree.plain_text(),
            blocks_json=tree.to_json())))
    out = tmp_path / "batch"
    paths, failures = batch_compile_each(ids, default_template(), str(out))
    assert len(paths) == 3 and failures == []
    for p in paths:
        assert Path(p).exists() and Path(p).stat().st_size > 5000


# ================================================================ F10 朗读
def test_tts_split_sentences():
    from gwtool.core.tts import split_sentences
    sents = split_sentences("第一句话。第二句！第三句？最后没有句号")
    assert len(sents) == 4
    assert sents[0] == "第一句话。"


def test_tts_availability_report():
    from gwtool.core.tts import available
    ok, desc = available()
    assert isinstance(ok, bool) and desc


# ================================================================ F16 便携
def test_portable_paths(tmp_path, monkeypatch):
    from gwtool import paths
    monkeypatch.setattr(paths, "_exe_base", lambda: tmp_path)
    paths.set_portable(True)
    d = paths.app_data_dir()
    assert d == tmp_path / "Data"
    paths.set_portable(False)
