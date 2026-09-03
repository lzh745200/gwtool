# -*- coding: utf-8 -*-
"""汇编生成 + 小册子 + PDF 渲染测试。"""
from pathlib import Path

import pytest

from gwtool.core.model import Block, DocTree
from gwtool.core.template import default_template
from gwtool.core import docxgen, booklet


def _sample_trees():
    t1 = DocTree(title="关于第一季度工作的报告")
    t1.blocks = [
        Block(type="heading", level=1, text="一、总体情况"),
        Block(type="paragraph", text="一季度各项工作平稳推进，现将有关情况报告如下。"),
        Block(type="heading", level=2, text="（一）主要成效"),
        Block(type="paragraph", text="重点任务按时完成，制度不断完善。"),
    ]
    t2 = DocTree(title="关于安全生产的通知")
    t2.blocks = [
        Block(type="heading", level=1, text="一、高度重视"),
        Block(type="paragraph", text="各单位要高度重视安全生产工作，压实责任。"),
    ]
    return [t1, t2]


def test_generate_docx_structure(tmp_path):
    tpl = default_template()
    tpl.toc_enabled = True
    out = tmp_path / "out.docx"
    docxgen.generate_docx(_sample_trees(), tpl, str(out))
    assert out.exists() and out.stat().st_size > 0

    from docx import Document
    from docx.oxml.ns import qn
    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    # 目录域存在
    assert any("TOC" in (p._p.xml or "") for p in doc.paragraphs)
    # 正文标题存在
    assert "一、总体情况" in texts
    # settings 包含 updateFields 与 evenAndOddHeaders
    xml = doc.settings.element.xml
    assert "updateFields" in xml
    assert "evenAndOddHeaders" in xml
    # 页脚包含 PAGE 域
    sec = doc.sections[0]
    assert "PAGE" in sec.footer.paragraphs[0]._p.xml
    assert "PAGE" in sec.even_page_footer.paragraphs[0]._p.xml
    # 正文行距 28 磅
    body_par = next(p for p in doc.paragraphs if p.text.startswith("一季度"))
    assert body_par.paragraph_format.line_spacing.pt == 28


def test_docx_reopenable_and_headings(tmp_path):
    """输出 docx 能被 python-docx 重新打开（OOXML 合法性），标题样式正确。"""
    out = tmp_path / "check.docx"
    docxgen.generate_docx(_sample_trees(), default_template(), str(out))
    from docx import Document
    doc = Document(str(out))
    styles = [p.style.name for p in doc.paragraphs if p.text == "一、总体情况"]
    assert styles and styles[0].startswith("Heading")


def test_booklet_order_math():
    # 8 页：8,1 | 2,7 | 6,3 | 4,5
    order = booklet.booklet_order(8)
    assert order == [[8, 1], [2, 7], [6, 3], [4, 5]]
    # 4 页
    assert booklet.booklet_order(4) == [[4, 1], [2, 3]]
    # 5 页 -> 补齐 8 页
    order5 = booklet.booklet_order(5)
    assert order5[0] == [8, 1] and len(order5) == 4
    # 12 页
    o12 = booklet.booklet_order(12)
    assert o12[0] == [12, 1] and o12[1] == [2, 11] and o12[3] == [4, 9]


def test_booklet_imposition(tmp_path, qapp):
    """先生成一份多页 A4 PDF，再重排为小册子，输出页数 = 补齐到4倍数后的输出页。"""
    from gwtool.core.pdfrender import trees_to_html, _render_html_to_pdf
    tpl = default_template()
    tpl.toc_enabled = False
    tpl.page_number_enabled = False
    trees = _sample_trees() + [DocTree(title="三", blocks=[Block(text="x" * 50)] * 60)]
    html = trees_to_html(trees, tpl, toc_pages=[])
    src = tmp_path / "src.pdf"
    _render_html_to_pdf(html, tpl, str(src))
    n_src = booklet.export_a4_pdf(str(src), str(src))
    assert n_src >= 4, f"源 PDF 仅 {n_src} 页"

    out = tmp_path / "booklet.pdf"
    n_out = booklet.make_booklet(str(src), str(out))
    assert n_out % 2 == 0 and n_out >= n_src // 2
    import pymupdf as fitz
    d = fitz.open(str(out))
    assert d.page_count == n_out
    # 两页并排：输出宽约为源页宽两倍
    w, h = d[0].rect.width, d[0].rect.height
    assert w > h
    d.close()


def test_pdf_render_full(tmp_path, qapp):
    """两遍渲染：目录页码真实填充 + 外侧页码盖章。"""
    from gwtool.core.pdfrender import render_compiled_pdf
    tpl = default_template()
    tpl.red_header.enabled = True
    tpl.toc_enabled = True
    out = tmp_path / "compiled.pdf"
    render_compiled_pdf(_sample_trees(), tpl, str(out))
    import pymupdf as fitz
    d = fitz.open(str(out))
    text_all = "\n".join(d[i].get_text() for i in range(d.page_count))
    d.close()
    # 页码格式存在
    assert "— 1 —" in text_all.replace(" ", " ") or "—1—" in text_all.replace(" ", "")
    # 目录标题出现
    assert "目" in text_all and "录" in text_all


def test_formatter_and_differ():
    from gwtool.core.formatter import run_full_cleanup
    from gwtool.core.differ import diff_to_html
    text = "一、标题1\n\n\n正文　　带多余空格 。\n2、3个数字混用\n（一）子项"
    new_text, log = run_full_cleanup(text)
    assert "多余空格" in "".join(log) or new_text != text
    html = diff_to_html("第一段内容。\n第二段。", "第一段修改后内容。\n第二段。\n新增段。")
    assert "ins" in html and "del" in html
