# -*- coding: utf-8 -*-
"""任意文档纠错增强：标记渲染、按块纠错、对话框交互回归。"""
from __future__ import annotations

import pytest

from gwtool.core import corrector
from gwtool.core.model import Block, DocTree, HEADING, PARAGRAPH, TABLE
from gwtool.db import dao


# ------------------------------------------------------------ 标记渲染
def test_to_marked_html_escapes_and_marks():
    text = '关于<布署>工作的通知 & 截止8月底'
    cs = corrector.check_text(text)
    assert cs, "样例应至少命中一处"
    html = corrector.to_marked_html(text, cs)
    assert "&lt;" in html and "&gt;" in html   # 原文的尖括号被转义
    assert 'text-decoration:underline' in html
    assert "〔" in html and "〕" in html    # 建议随附
    assert html.count("<a name=") == len(cs)


def test_to_marked_html_anchor_prefix_and_no_suggestion():
    cs = corrector.check_text("工作布署已经完成")
    html = corrector.to_marked_html("工作布署已经完成", cs,
                                    show_suggestion=False, anchor_prefix="x")
    assert '<a name="x0"></a>' in html
    assert "〔" not in html


def test_to_marked_html_overlap_safe():
    cs = corrector.check_text("工作布署已经完成")
    doubled = cs + cs                      # 人为重叠
    html = corrector.to_marked_html("工作布署已经完成", doubled)
    assert html.count("<a name=") == len(cs)   # 重叠不重复渲染


def test_paragraph_no():
    text = "第一段\n第二段\n第三段"
    assert corrector.paragraph_no(text, 0) == 1
    assert corrector.paragraph_no(text, 5) == 2
    assert corrector.paragraph_no(text, 999) == 3


# ------------------------------------------------------------ 按块纠错
def test_correct_block_applies_and_skips():
    text = "工作布署已经完成，报名截止8月底。"
    fixed, cs = corrector.correct_block(text)
    assert "工作部署已经完成" in fixed
    assert "截至8月底" in fixed
    # 显式包含数字用法时也不丢
    fixed2, cs2 = corrector.correct_block(text, skip_categories=())
    assert len(cs2) >= len(cs)


def test_correct_block_on_table_cell():
    fixed, cs = corrector.correct_block("项目布署情况")
    assert fixed == "项目部署情况" and cs


# ------------------------------------------------------------ 树块转换
def test_tree_to_blocks_and_back(tmp_db):
    tree = DocTree(title="布署通知", blocks=[
        Block(type=HEADING, level=1, text="一、总体要求"),
        Block(type=PARAGRAPH, text="工作布署截止本月底。"),
        Block(type=TABLE, rows=[["项目", "布署情况"], ["甲", "截止8月底"]]),
    ])
    blocks = corrector.tree_to_blocks(tree)
    assert [b["kind"] for b in blocks] == ["heading", "para", "table"]  # 标题不折块
    # 逐块纠错：标题单独、正文、单元格
    title = corrector.correct_block(tree.title)[0]
    for b in blocks:
        if b["kind"] == "table":
            b["rows"] = [[corrector.correct_block(c)[0] for c in row] for row in b["rows"]]
        else:
            b["text"] = corrector.correct_block(b["text"])[0]
    new_tree = corrector.blocks_to_tree(title, blocks)
    assert new_tree.title == "部署通知"
    body = "\n".join(b.text for b in new_tree.blocks if b.type == PARAGRAPH)
    assert "工作部署截至本月底" in body
    tbl = next(b for b in new_tree.blocks if b.type == TABLE)
    assert tbl.rows[0] == ["项目", "部署情况"] and tbl.rows[1] == ["甲", "截至8月底"]


# ------------------------------------------------------------ 对话框交互
def _drain(dlg):
    """等后台扫描线程结束再断言/收尾，避免 QThread 销毁崩溃。"""
    if dlg._worker is not None and dlg._worker.isRunning():
        dlg._worker.wait(15000)


def test_dialog_full_flow(tmp_db, qapp, monkeypatch):
    from PySide6.QtWidgets import QDialog, QMessageBox
    monkeypatch.setattr(QDialog, "exec", lambda self: QDialog.DialogCode.Rejected)
    monkeypatch.setattr(QMessageBox, "information", staticmethod(lambda *a, **k: None))
    monkeypatch.setattr(QMessageBox, "warning", staticmethod(lambda *a, **k: None))
    monkeypatch.setattr(QMessageBox, "critical", staticmethod(lambda *a, **k: None))
    monkeypatch.setattr(QMessageBox, "question", staticmethod(
        lambda *a, **k: QMessageBox.StandardButton.No))

    from gwtool.ui.correct_dialog import AnyDocCorrectDialog, _scan_blocks

    dlg = AnyDocCorrectDialog()
    text = "关于布署工作的通知\n工作布署截止本月底。"
    dlg._load_text(text, "探测")

    # 工作线程函数同步执行，结果回填（模拟 ok 信号）
    _drain(dlg)
    dlg._on_checked(_scan_blocks(dlg._blocks))
    n = dlg.listw.count()
    assert n >= 2, "两段至少各命中一处"

    # 标记视图含锚点与高亮
    html = dlg.browser.toHtml()
    assert "<a name=" in html

    # 选中第一条 → 应用 → 文本被修正
    dlg.listw.setCurrentRow(0)
    dlg.apply_selected()
    assert "部署" in "\n".join(b["text"] for b in dlg._blocks)

    # 全部应用（排除数字用法）后重扫：不再有高置信错别字
    dlg.apply_all_blocks()
    _drain(dlg)
    dlg._on_checked(_scan_blocks(dlg._blocks))
    remaining_wrong = {c.wrong for cs in dlg._corrs for c in cs
                       if c.category in ("错别字", "用户词库")}
    assert not remaining_wrong, str(remaining_wrong)

    # 导出 TXT：内容为修正后文本
    import tempfile
    from pathlib import Path
    out = Path(tempfile.mkdtemp()) / "fixed.txt"
    monkeypatch.setattr(
        "gwtool.ui.correct_dialog.QFileDialog.getSaveFileName",
        staticmethod(lambda *a, **k: (str(out), "")))
    dlg.export_txt()
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "部署" in content and "布署" not in content

    # 忽略此词：写库并重扫后不再提示
    dlg._corrs = _scan_blocks(dlg._blocks)
    dlg._render()
    if dlg.listw.count():
        dlg.listw.setCurrentRow(0)
        word = dlg._current_hit()["corr"].wrong
        dlg.ignore_selected()
        assert word in dao.all_ignore_words()
    _drain(dlg)


def test_dialog_file_mode_preserves_structure(tmp_db, qapp, monkeypatch):
    """文件来源：标题/表格结构经树转换保留，修正后导出 DOCX 可重开。"""
    from PySide6.QtWidgets import QDialog, QMessageBox
    monkeypatch.setattr(QDialog, "exec", lambda self: QDialog.DialogCode.Rejected)
    monkeypatch.setattr(QMessageBox, "warning", staticmethod(lambda *a, **k: None))
    monkeypatch.setattr(QMessageBox, "information", staticmethod(lambda *a, **k: None))
    monkeypatch.setattr(QMessageBox, "critical", staticmethod(lambda *a, **k: None))

    import tempfile
    from pathlib import Path
    from docx import Document as DX
    from docx.shared import Pt
    src = Path(tempfile.mkdtemp()) / "样例.docx"
    d = DX()
    d.add_heading("布署方案", level=1)
    d.add_paragraph("工作布署如下。")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "项目"
    t.cell(0, 1).text = "布署情况"
    t.cell(1, 0).text = "甲"
    t.cell(1, 1).text = "截止8月底"
    d.save(str(src))

    from gwtool.ui.correct_dialog import AnyDocCorrectDialog, _scan_blocks
    dlg = AnyDocCorrectDialog()
    monkeypatch.setattr(
        "gwtool.ui.correct_dialog.QFileDialog.getOpenFileName",
        staticmethod(lambda *a, **k: (str(src), "")))
    dlg.pick_file()
    kinds = [b["kind"] for b in dlg._blocks]
    assert "heading" in kinds and "table" in kinds

    _drain(dlg)
    dlg._on_checked(_scan_blocks(dlg._blocks))
    # 全部应用（含表格单元格）
    dlg.apply_all_blocks()

    out = Path(tempfile.mkdtemp()) / "fixed.docx"
    monkeypatch.setattr(
        "gwtool.ui.correct_dialog.QFileDialog.getSaveFileName",
        staticmethod(lambda *a, **k: (str(out), "")))
    dlg.export_docx()
    assert out.exists() and out.stat().st_size > 5000
    dd = DX(str(out))
    cells = [c.text for tb in dd.tables for row in tb.rows for c in row.cells]
    assert "部署情况" in cells and "截至8月底" in cells
    _drain(dlg)
