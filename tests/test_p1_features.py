# -*- coding: utf-8 -*-
"""P1 回归：docx 表格、目录分页、材料编号、批量容错、备份模板回读、快照级联。"""
from gwtool.core.model import Block, DocTree, HEADING, PARAGRAPH, TABLE


def test_docx_table_roundtrip(tmp_path):
    """docx 中的表格须解析为 table 块，并在生成端还原为 Word 表格。"""
    from docx import Document as NewDocx
    src = tmp_path / "t.docx"
    doc = NewDocx()
    doc.add_paragraph("标题一", style="Heading 1")
    table = doc.add_table(rows=2, cols=3)
    cells = (("a1", "b1", "c1"), ("a2", "b2", "c2"))
    for i, row in enumerate(cells):
        for j, v in enumerate(row):
            table.cell(i, j).text = v
    doc.add_paragraph("结尾段")
    doc.save(str(src))

    from gwtool.core.parsers.docx_parser import parse_docx
    tree = parse_docx(str(src))
    tables = [b for b in tree.blocks if b.type == TABLE]
    assert len(tables) == 1
    assert tables[0].rows == [list(r) for r in cells]
    # 表格前后的段落仍在正确顺序上
    kinds = [b.type for b in tree.blocks]
    assert kinds[0] == HEADING and kinds[-1] == PARAGRAPH

    from gwtool.core.docxgen import generate_docx
    from gwtool.core.template import default_template
    tpl = default_template()
    out = tmp_path / "out.docx"
    generate_docx([tree], tpl, str(out))
    check = NewDocx(str(out))
    assert len(check.tables) >= 1
    got = [c.text for c in check.tables[0].rows[0].cells]
    assert got == ["a1", "b1", "c1"]


def test_toc_chunks_math():
    """目录分页切片：总数守恒、首页留标题空间。"""
    from gwtool.core.pdfrender import _TOC_LINES_PER_PAGE, _toc_chunks
    assert _toc_chunks(0) == []
    assert _toc_chunks(5) == [5]
    assert _toc_chunks(20) == [20]
    assert _toc_chunks(22) == [20, 2]
    assert _toc_chunks(50) == [20, 22, 8]
    for n in (0, 1, 20, 21, 43, 100, 500):
        assert sum(_toc_chunks(n)) == n
    assert max(_toc_chunks(200)) <= _TOC_LINES_PER_PAGE


def test_trees_to_html_toc_paged(tmp_db, qapp):
    """长目录须拆分多页（存在额外分页符）。"""
    from gwtool.core.pdfrender import trees_to_html
    from gwtool.core.template import default_template
    trees = []
    for i in range(30):
        t = DocTree(title=f"材料{i}")
        t.blocks.append(Block(type=HEADING, level=1, text=f"标题{i}"))
        trees.append(t)
    tpl = default_template()
    toc = [[(f"标题{i}", 1, 0) for i in range(30)]]
    html = trees_to_html(trees, tpl, toc_pages=toc, toc_placeholder=True)
    # 30 行 -> 两页目录 -> 目录内部 1 个额外分页符
    assert html.count("page-break-before") >= 1


def test_material_label():
    from gwtool.core.template import material_label
    assert material_label("", 1) == ""
    assert material_label("材料{n}：", 1) == "材料一："
    assert material_label("材料{n}：", 2) == "材料二："
    assert material_label("材料{n}：", 12) == "材料十二："
    assert material_label("附件", 3) == "附件"


def test_material_prefix_in_docx(tmp_path):
    """材料编号前缀落到生成的 docx 标题。"""
    from docx import Document as NewDocx
    from gwtool.core.docxgen import generate_docx
    from gwtool.core.template import DocTemplate
    t = DocTemplate()
    t.toc_enabled = False
    t.material_title_prefix = "材料{n}："
    tree = DocTree(title="某某实施方案")
    tree.blocks.append(Block(text="正文内容"))
    out = tmp_path / "m.docx"
    generate_docx([tree], t, str(out))
    doc = NewDocx(str(out))
    assert any(p.text == "材料一：某某实施方案" for p in doc.paragraphs)


def test_batch_compile_failures(tmp_db, tmp_path, monkeypatch):
    """批量汇编：单份失败不中断整批，失败清单逐条返回。"""
    from gwtool import app as appmod
    from gwtool.core import batch
    from gwtool.core.template import default_template
    from gwtool.db import dao

    appmod.ensure_database_seeded()
    ids = [dao.add_document(dao.Document(title=f"d{i}", content_text=f"内容{i}"))
           for i in range(3)]
    real = batch.docxgen.generate_docx

    def fake(trees, tpl, out):
        if "d1.docx" in str(out):
            raise RuntimeError("boom")
        return real(trees, tpl, out)

    monkeypatch.setattr(batch.docxgen, "generate_docx", fake)
    progress = []

    paths, failures = batch.batch_compile_each(
        ids, default_template(), str(tmp_path),
        progress_cb=lambda i, n: progress.append((i, n)))
    assert len(paths) == 2
    assert len(failures) == 1 and failures[0][0] == "d1"
    assert progress[-1] == (3, 3)   # 进度与实际完成数同步


def test_backup_restore_templates(tmp_db):
    """备份包内 templates.json 恢复时回写模板表（旧版只写不读）。"""
    from gwtool import app as appmod
    from gwtool.core import backup
    from gwtool.db import dao

    appmod.ensure_database_seeded()
    dao.save_template("自定义模板A", '{"name":"自定义模板A"}', is_default=False)
    path = backup.create_backup(note="t")
    for t in dao.list_templates():
        if t["name"] == "自定义模板A":
            dao.delete_template(t["id"])
    assert all(t["name"] != "自定义模板A" for t in dao.list_templates())
    backup.restore_backup(path)
    assert any(t["name"] == "自定义模板A" for t in dao.list_templates())


def test_delete_document_cascades_snapshots(tmp_db):
    """删除文档须级联清理其历史快照。"""
    from gwtool.db import dao
    did = dao.add_document(dao.Document(title="x", content_text="内容"))
    dao.add_snapshot(did, "x", "内容", reason="auto")
    dao.add_snapshot(did, "x", "内容2", reason="auto")
    assert len(dao.list_snapshots(did)) == 2
    dao.delete_document(did)
    assert dao.list_snapshots(did) == []


def test_backup_wal_checkpoint(tmp_db):
    """备份前执行 WAL checkpoint（Truncate 模式下 WAL 文件应清零）。"""
    from gwtool import app as appmod
    from gwtool.core import backup
    from gwtool.db import dao, connection as dbconn

    appmod.ensure_database_seeded()
    dao.add_document(dao.Document(title="wal", content_text="写入内容"))
    path = backup.create_backup(note="wal")
    wal = dbconn.current_db_file().with_name("gwtool.db-wal")
    if wal.exists():
        assert wal.stat().st_size == 0, "checkpoint(TRUNCATE) 后 WAL 应为空"
    assert backup.list_backups()
