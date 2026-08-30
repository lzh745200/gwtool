# -*- coding: utf-8 -*-
"""端到端自检脚本：在真实数据目录上走通「种子导入→建分类→导入材料→
汇编 docx→A4 PDF→A3 小册子→检索→写作参考→纠错→备份」全流程。

用法：.venv/Scripts/python scripts/e2e_check.py
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import os

if not sys.platform.startswith("win"):
    os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")


def main() -> int:
    ok = 0
    fail = []

    def step(name, cond, detail=""):
        nonlocal ok
        if cond:
            ok += 1
            print(f"  [PASS] {name} {detail}")
        else:
            fail.append(name)
            print(f"  [FAIL] {name} {detail}")

    # 0) 数据库（真实目录）
    from gwtool.db import connection as dbconn
    from gwtool import app
    dbconn.configure(__import__("gwtool.paths", fromlist=["db_path"]).db_path())
    t0 = __import__("time").time()
    app.ensure_database_seeded()
    seed_t = __import__("time").time() - t0
    from gwtool.db import dao
    step("首启动种子导入(秒)", seed_t < 15, f"耗时 {seed_t:.1f}s")
    step("纠错库≥3万", dao.count_error_pairs() >= 30000,
         f"{dao.count_error_pairs()} 条")

    # 1) 准备样例材料
    from docx import Document as DX
    sample_dir = Path(tempfile.mkdtemp(prefix="gwtool_e2e_"))
    p1 = sample_dir / "关于季度工作报告.docx"
    d = DX()
    d.add_heading("关于第一季度工作的报告", level=1)
    d.add_paragraph("一季度以来，各项工作平稳有序推进，重点任务完成情况良好。")
    d.add_heading("（一）主要成效", level=2)
    d.add_paragraph("项目建设布署已经完成，资金拨付截止本月底。")   # 故意放错误
    d.save(str(p1))
    p2 = sample_dir / "会议纪要.txt"
    p2.write_text("关于安全生产的会议纪要\n会议强调，要压实安全生产责任，坚决防范各类事故发生。",
                  encoding="gbk")
    p3 = sample_dir / "政策要点.md"
    p3.write_text("# 乡村振兴政策要点\n\n坚持农业农村优先发展，巩固拓展脱贫攻坚成果。\n",
                  encoding="utf-8")

    # 2) 分类 + 导入（幂等：重复内容复用已入库文档）
    existing = {d.title: d.id for d in dao.list_documents()}
    cat_name = "E2E自检分类"
    cat = next((c.id for c in dao.list_categories() if c.name == cat_name), None)
    if cat is None:
        cat = dao.add_category(cat_name)
    from gwtool.core.importer import parse_any
    ids = []
    for p in (p1, p2, p3):
        r = parse_any(str(p))
        step(f"解析 {p.name}", r.ok, r.error if not r.ok else "")
        if r.ok:
            did = dao.add_document(dao.Document(
                title=r.tree.title, content_text=r.tree.plain_text(),
                blocks_json=r.tree.to_json(), file_type=p.suffix.lstrip("."),
                category_id=cat))
            if did < 0:
                did = existing.get(r.tree.title, 0)
            if did:
                ids.append(did)

    # 3) 汇编 docx
    from gwtool.core.template import default_template
    from gwtool.core.compiler import compile_docx, CompileRequest
    tpl = default_template()
    tpl.red_header.org = "××市自主创新示范区政府办公室文件"
    tpl.red_header.doc_number = "×政办发〔2026〕12号"
    out_docx = sample_dir / "汇编成果.docx"
    compile_docx(CompileRequest(doc_ids=ids, template=tpl, out_docx=str(out_docx)))
    step("汇编 docx", out_docx.exists() and out_docx.stat().st_size > 10000,
         f"{out_docx.stat().st_size/1024:.0f} KB")
    from docx import Document as DX2
    dd = DX2(str(out_docx))
    xml = dd.settings.element.xml
    step("docx 目录域+奇偶页脚", "TOC" in "".join(p._p.xml for p in dd.paragraphs)
         and "updateFields" in xml and "evenAndOddHeaders" in xml)

    # 4) 内置渲染 A4 PDF（含目录页码与页码盖章）
    from PySide6.QtWidgets import QApplication
    _qapp = QApplication.instance() or QApplication([])
    from gwtool.core import pdfrender
    trees = __import__("gwtool.core.compiler", fromlist=["load_trees"]).load_trees(ids, [])
    out_pdf = sample_dir / "汇编成果.pdf"
    pdfrender.render_compiled_pdf(trees, tpl, str(out_pdf))
    step("A4 PDF(两遍渲染)", out_pdf.exists() and out_pdf.stat().st_size > 10000,
         f"{out_pdf.stat().st_size/1024:.0f} KB")

    # 5) A3 骑马钉小册子
    from gwtool.core.booklet import make_booklet
    out_bk = sample_dir / "汇编成果_小册子A3.pdf"
    n = make_booklet(str(out_pdf), str(out_bk))
    step("A3 小册子", out_bk.exists() and n > 0, f"{n} 页")

    # 6) 全文检索
    res = dao.search_documents("安全生产")
    step("FTS5 检索『安全生产』", len(res) >= 1, f"{len(res)} 条")

    # 7) 写作参考
    from gwtool.core import reference
    items = reference.lookup("发展")
    step("写作参考检索", len(items) > 0, f"{len(items)} 条")

    # 8) 纠错（验收对）
    from gwtool.core import corrector
    corr = corrector.check_text("工作布署已完成，报名截止本月底。")
    pairs = {c.wrong: c.suggestion for c in corr}
    step("纠错：布署→部署", pairs.get("布署") == "部署")
    step("纠错：截止→截至", pairs.get("截止") == "截至")

    # 9) 备份
    from gwtool.core.backup import create_backup
    bz = create_backup(note="E2E自检")
    step("一键备份", Path(bz).exists(), str(Path(bz).name))

    print(f"\n===== E2E 自检结果：{ok} 项通过，{len(fail)} 项失败 =====")
    if fail:
        print("失败项：", fail)
    print(f"样例输出目录：{sample_dir}")
    return 0 if not fail else 1


if __name__ == "__main__":
    sys.exit(main())
