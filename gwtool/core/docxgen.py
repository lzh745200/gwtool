# -*- coding: utf-8 -*-
"""公文 DOCX 生成引擎（python-docx）。

生成要素：封面 -> 红头 -> 目录(TOC域, 打开自动更新页码) -> 正文 -> 版记；
格式要点（GB/T 9704 常用实践）：
  - 正文仿宋_GB2312 三号 16pt、固定行距28磅、首行缩进2字符(w:firstLineChars=200)；
  - 页码宋体四号、奇数页右对齐/偶数页左对齐（外侧），格式 “— 1 —”；
  - 标题用 Heading 1-3 样式（供 TOC 抓取），字体覆盖为黑体/楷体/仿宋；
  - 全部使用标准 OOXML 元素，WPS/永中 Office 可正常打开不跑版。
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Mm, Pt, RGBColor

from .model import DocTree, HEADING, LIST_ITEM, PARAGRAPH
from .template import DocTemplate, HeadingStyle

_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# ------------------------------------------------------------------ 底层工具
def _set_run_font(run, cn_font: str, size_pt: float, bold: bool = False,
                  color: tuple[int, int, int] | None = None) -> None:
    """同时设置西文与中文（eastAsia）字体，缺一不可。"""
    run.font.name = cn_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), cn_font)
    rfonts.set(qn("w:hAnsi"), cn_font)
    rfonts.set(qn("w:eastAsia"), cn_font)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_line_exact(par, pt: float) -> None:
    pf = par.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(pt)


def _set_first_line_chars(par, chars: float) -> None:
    """以“字符”为单位的首行缩进（w:firstLineChars=200 即 2 字符）。"""
    ppr = par._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    ind.set(qn("w:firstLineChars"), str(int(chars * 100)))
    ind.set(qn("w:firstLine"), str(int(chars * 240)))  # 兜底twips(12pt字面值)


def _bottom_border(par, color: str = "FF0000", sz: int = 15) -> None:
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))       # 1/8 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.append(pbdr)


def _add_field(par, instr: str, hint: str = "") -> None:
    """插入 Word 域（PAGE / TOC 等）。"""
    r1 = par.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fld.set(qn("w:dirty"), "true")   # 打开文档时重算
    r1._element.append(fld)
    r2 = par.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._element.append(it)
    r3 = par.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r3._element.append(sep)
    if hint:
        par.add_run(hint)
    r4 = par.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r4._element.append(end)


def _enable_update_fields(doc) -> None:
    """settings.xml 加 w:updateFields，Word/WPS 打开时自动更新域（目录页码）。"""
    settings = doc.settings.element
    upd = settings.find(qn("w:updateFields"))
    if upd is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


def _enable_even_odd_footers(doc) -> None:
    settings = doc.settings.element
    eo = settings.find(qn("w:evenAndOddHeaders"))
    if eo is None:
        eo = OxmlElement("w:evenAndOddHeaders")
        settings.append(eo)


# ------------------------------------------------------------------ 主流程
def generate_docx(trees: list[DocTree], tpl: DocTemplate, out_path: str) -> str:
    doc = Document()

    # 页面设置
    sec = doc.sections[0]
    sec.page_width = Mm(tpl.page_width_mm)
    sec.page_height = Mm(tpl.page_height_mm)
    sec.top_margin = Mm(tpl.margin_top_mm)
    sec.bottom_margin = Mm(tpl.margin_bottom_mm)
    sec.left_margin = Mm(tpl.margin_left_mm)
    sec.right_margin = Mm(tpl.margin_right_mm)
    sec.footer_distance = Mm(20)

    # 页码：奇偶页脚（外侧）+ 打开自动更新域
    _enable_even_odd_footers(doc)
    _enable_update_fields(doc)
    if tpl.page_number_enabled:
        _build_footers(sec, tpl)

    # 1) 封面
    if tpl.cover.enabled:
        _build_cover(doc, tpl)
        doc.add_page_break()

    # 2) 红头
    if tpl.red_header.enabled:
        _build_red_header(doc, tpl)

    # 3) 目录（TOC 域）
    if tpl.toc_enabled:
        if tpl.red_header.enabled:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(tpl.toc_title)
        _set_run_font(run, "宋体", 22, bold=True)
        toc_p = doc.add_paragraph()
        _add_field(toc_p, r' TOC \o "1-3" \h \z \u ', "（打开文档后按提示更新目录页码）")
        doc.add_page_break()

    # 4) 正文
    for tree in trees:
        if tpl.insert_material_titles and tree.title:
            _add_heading(doc, tree.title, 1, tpl.h1, tpl.line_spacing_pt)
        for blk in tree.effective_blocks(tpl.insert_material_titles):
            if blk.type == HEADING:
                hs = {1: tpl.h1, 2: tpl.h2, 3: tpl.h3}.get(blk.level, tpl.h4)
                _add_heading(doc, blk.text, min(blk.level, 3), hs, tpl.line_spacing_pt)
            elif blk.type == LIST_ITEM:
                _add_body(doc, blk.text, tpl, indent=tpl.first_line_indent_chars)
            else:
                _add_body(doc, blk.text, tpl)

    # 5) 版记
    if tpl.colophon.enabled and tpl.colophon.lines:
        doc.add_page_break()
        _build_colophon(doc, tpl)

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))

    # 水印（页眉 VML 艺术字）
    if tpl.watermark_text.strip():
        from .watermark import add_watermark_docx
        add_watermark_docx(str(out), tpl.watermark_text.strip())
    return str(out)


# ------------------------------------------------------------------ 部件
def _build_footers(sec, tpl: DocTemplate) -> None:
    """奇数页页码右对齐、偶数页左对齐（外侧），宋体四号。"""
    def make(footer, align):
        footer.is_linked_to_previous = False
        par = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        par.text = ""
        par.alignment = align
        _set_line_exact(par, 20)
        pre, post = "", ""
        fmt = tpl.page_number_format or "— {page} —"
        if "{page}" in fmt:
            pre, post = fmt.split("{page}", 1)
        if pre:
            r = par.add_run(pre)
            _set_run_font(r, tpl.page_number_font, tpl.page_number_size_pt)
        _add_field(par, r" PAGE \* MERGEFORMAT ", "1")
        if post:
            r = par.add_run(post)
            _set_run_font(r, tpl.page_number_font, tpl.page_number_size_pt)
        # 域内提示字号
        for rr in par.runs:
            _set_run_font(rr, tpl.page_number_font, tpl.page_number_size_pt)

    make(sec.footer, WD_ALIGN_PARAGRAPH.RIGHT)     # 奇数页（默认页脚）
    make(sec.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT)


def _build_cover(doc, tpl: DocTemplate) -> None:
    cover = tpl.cover
    for _ in range(4):
        _spacer(doc, tpl)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_exact(p, 40)
    run = p.add_run(cover.title or tpl.name)
    _set_run_font(run, "方正小标宋简体", 30, color=(0, 0, 0))
    if cover.subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_exact(p2, 32)
        r2 = p2.add_run(cover.subtitle)
        _set_run_font(r2, "楷体_GB2312", 18)
    for _ in range(4):
        _spacer(doc, tpl)
    for line in [cover.org, cover.date, *cover.extra_lines]:
        if not line:
            continue
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_exact(p3, 28)
        r3 = p3.add_run(line)
        _set_run_font(r3, "仿宋_GB2312", 16)


def _build_red_header(doc, tpl: DocTemplate) -> None:
    rh = tpl.red_header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_line_exact(p, rh.org_size_pt * 1.35)
    run = p.add_run(rh.org)
    _set_run_font(run, rh.org_font if rh.org_font != "方正小标宋简体" else "方正小标宋简体",
                  rh.org_size_pt, color=(255, 0, 0))
    if rh.doc_number:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_line_exact(p2, 28)
        r2 = p2.add_run(rh.doc_number)
        _set_run_font(r2, rh.doc_number_font, rh.doc_number_size_pt)
    if rh.red_line:
        pl = doc.add_paragraph()
        _set_line_exact(pl, 12)
        _bottom_border(pl, color="FF0000", sz=15)


def _add_heading(doc, text: str, level: int, hs: HeadingStyle, line_pt: float) -> None:
    """用 Heading 样式（供目录抓取）并覆盖字体。"""
    par = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    par.alignment = _ALIGN.get(hs.align, WD_ALIGN_PARAGRAPH.LEFT)
    _set_line_exact(par, line_pt)
    if hs.indent_chars:
        _set_first_line_chars(par, hs.indent_chars)
    run = par.add_run(text)
    _set_run_font(run, hs.font, hs.size_pt, bold=hs.bold, color=(0, 0, 0))


def _add_body(doc, text: str, tpl: DocTemplate, indent: float | None = None) -> None:
    par = doc.add_paragraph()
    par.alignment = _ALIGN.get(tpl.align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _set_line_exact(par, tpl.line_spacing_pt)
    ind = tpl.first_line_indent_chars if indent is None else indent
    if ind:
        _set_first_line_chars(par, ind)
    run = par.add_run(text)
    _set_run_font(run, tpl.body_font, tpl.body_size_pt)


def _build_colophon(doc, tpl: DocTemplate) -> None:
    """版记：横线分隔的条目区，置于末页。"""
    lines = tpl.colophon.lines
    for i, line in enumerate(lines):
        par = doc.add_paragraph()
        _set_line_exact(par, 24)
        if i == 0:
            _top_and_bottom(par)
        elif i == len(lines) - 1:
            ppr = par._p.get_or_add_pPr()
            pbdr = OxmlElement("w:pBdr")
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "12")
            b.set(qn("w:space"), "1")
            b.set(qn("w:color"), "000000")
            pbdr.append(b)
            ppr.append(pbdr)
        run = par.add_run(line)
        _set_run_font(run, "仿宋_GB2312", 14)


def _top_and_bottom(par) -> None:
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for tag in ("w:top", "w:bottom"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "1")
        el.set(qn("w:color"), "000000")
        pbdr.append(el)
    ppr.append(pbdr)


def _spacer(doc, tpl: DocTemplate) -> None:
    p = doc.add_paragraph()
    _set_line_exact(p, tpl.line_spacing_pt)
