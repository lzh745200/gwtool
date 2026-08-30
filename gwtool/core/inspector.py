# -*- coding: utf-8 -*-
"""GB/T 9704 公文格式体检。

两级检查：
  - inspect_text：纯文本级（标题编号链条、成文日期、发文字号、结束语与文种匹配等）
  - inspect_docx：docx 级（页边距、正文字体字号行距、标题字体、页码域）
输出 Finding 列表：severity = error/warn/info。
"""
from __future__ import annotations

import re
from dataclasses import dataclass

_GB = {
    "margin": (37, 35, 28, 26),          # 上 下 左 右 mm
    "body_font": {"仿宋_GB2312", "仿宋", "FangSong", "仿宋_GB2312-0"},
    "h1_font": {"黑体", "SimHei", "Heiti"},
    "h2_font": {"楷体_GB2312", "楷体", "KaiTi"},
    "body_size": 16.0,                    # 三号
    "line_spacing": 28.0,                 # 磅
}

_CHAIN = [  # 层级顺序
    (re.compile(r"^[一二三四五六七八九十]+、"), "一、"),
    (re.compile(r"^[（(][一二三四五六七八九十]+[)）]"), "（一）"),
    (re.compile(r"^\d{1,2}[..]"), "1."),
    (re.compile(r"^[（(]\d{1,2}[)）]"), "（1）"),
]

_CN_ORD = "一二三四五六七八九十"


@dataclass
class Finding:
    severity: str   # error / warn / info
    item: str       # 检查项
    detail: str     # 说明（含位置）

    @property
    def label(self) -> str:
        icon = {"error": "✖", "warn": "▲", "info": "ℹ"}.get(self.severity, "·")
        return f"{icon} [{self.item}] {self.detail}"


def _classify(line: str) -> tuple[int, str] | None:
    for lv, (rx, _) in enumerate(_CHAIN):
        if rx.match(line):
            return lv, line
    return None


def _cn_to_int(s: str) -> int:
    """中文数字 -> 整数（一~九十九，覆盖公文编号范围）。

    注意“一”必须映射为 1（用 零一二… 映射表，不能用下标 0 的集合）。
    """
    digits = "零一二三四五六七八九"
    total, cur = 0, 0
    for ch in s:
        if ch == "十":
            total += (cur or 1) * 10
            cur = 0
        elif ch == "百":
            total = (total + (cur or 1)) * 100
            cur = 0
        elif ch in digits:
            cur = cur * 10 + digits.index(ch)
    return total + cur


def inspect_text(text: str, kind_hint: str = "") -> list[Finding]:
    out: list[Finding] = []
    lines = [ln.strip() for ln in text.split("\n")]
    title = next((ln for ln in lines if ln), "")
    kind = kind_hint or next(
        (k for k in ("请示", "报告", "通知", "函", "批复", "通报", "意见", "纪要")
         if k in title), "")

    # ---- 标题编号链条 ----
    counters = [0, 0, 0, 0]
    seen_kinds: list[int] = []
    for ln_no, ln in enumerate(lines, 1):
        cls = _classify(ln)
        if not cls:
            continue
        lv, line = cls
        body = _CHAIN[lv][0].match(line).group(0)
        if lv == 0:
            num = _cn_to_int(re.sub(r"[、.．]", "", body))
        elif lv == 1:
            num = _cn_to_int(re.sub(r"[（）()、.．\s]", "", body))
        else:
            num = int(re.sub(r"\D", "", body) or 0)
        expect = counters[lv] + 1
        if num != expect:
            if num > expect:
                out.append(Finding("warn", "标题编号",
                                   f"第{ln_no}行「{body}…」跳号，应为第{expect}项"))
            else:
                out.append(Finding("warn", "标题编号",
                                   f"第{ln_no}行「{body}…」编号重复或回退"))
        counters[lv] = num
        counters[lv + 1:] = [0] * (3 - lv)
        if seen_kinds and lv > max(seen_kinds) + 1:
            out.append(Finding("warn", "标题层级",
                               f"第{ln_no}行从 {_CHAIN[max(seen_kinds)][1]} "
                               f"直接跳到 {_CHAIN[lv][1]}，层级不连贯"))
        if lv not in seen_kinds:
            seen_kinds.append(lv)
        if lv < seen_kinds[-1] and lv + 1 < len(counters):
            pass

    # ---- 发文字号 ----
    for m in re.finditer(r"[〔\[（(]\s*(\d{4})\s*[〕\]）)]", text):
        bracket = text[m.start()]
        if bracket in "[(":
            out.append(Finding("error", "发文字号",
                               f"「{m.group(0)}」应使用六角括号〔〕，如“×政发〔{m.group(1)}〕5号”"))

    # ---- 成文日期 ----
    for m in re.finditer(r"[二〇○○○O零一二三四五六七八九]{4}年", text):
        out.append(Finding("info", "成文日期",
                           f"「{m.group(0)}」建议改为阿拉伯数字，如“2026年”"))
    for m in re.finditer(r"(\d{4})年(\d{1,2})月(\d{1,2})日", text):
        if len(m.group(2)) == 2 and m.group(2)[0] == "0" \
                or len(m.group(3)) == 2 and m.group(3)[0] == "0":
            out.append(Finding("warn", "成文日期",
                               f"「{m.group(0)}」月/日不应有前导零，应为“{int(m.group(2))}月{int(m.group(3))}日”"))
        break
    for m in re.finditer(r"\d{4}[./]\d{1,2}[./]\d{1,2}", text):
        out.append(Finding("warn", "成文日期",
                           f"「{m.group(0)}」公文成文日期建议写为“2026年8月30日”式"))
        break

    # ---- 结束语与文种匹配 ----
    closings = {
        "请示": ("妥否，请批示", "以上请示如无不妥，请批复", "请批示"),
        "报告": ("特此报告", "专此报告"),
        "通知": ("特此通知",),
        "通报": ("特此通报",),
        "函": ("为盼", "为荷", "函复"),
        "批复": ("请遵照执行",),
    }
    if kind in closings:
        if not any(c in text for c in closings[kind]):
            out.append(Finding("warn", "结束语",
                               f"文种「{kind}」通常应有结束语，如“{closings[kind][0]}。”"))
    if kind == "报告" and re.search(r"请批示|请批复", text):
        out.append(Finding("error", "文种混用",
                           "“报告”不得夹带请示事项（出现“请批示/请批复”字样）"))
    if kind == "请示" and "一文" not in text and text.count("请示事项") > 1:
        out.append(Finding("info", "请示规则", "请示应一文一事，请核对是否合并多事项"))

    # ---- 主送机关 ----
    if kind in ("通知", "通报", "请示", "报告", "函", "批复", "意见"):
        has_recipient = any(
            ln.endswith(("：", ":")) and len(ln) <= 40 and i <= 6
            for i, ln in enumerate(lines[:8]))
        if not has_recipient:
            out.append(Finding("warn", "主送机关",
                               f"「{kind}」应顶格标注主送机关（如“各科室：”）"))

    # ---- 篇幅参考 ----
    n_chars = len(re.sub(r"\s", "", text))
    if kind and n_chars > 3000:
        out.append(Finding("info", "篇幅",
                           f"全文约 {n_chars} 字，一般性「{kind}」建议精炼篇幅"))

    out.extend(_check_classification(lines))
    out.extend(_check_attachments(lines))
    return out


def _check_classification(lines: list[str]) -> list[Finding]:
    """密级标注：绝密/机密/秘密 开头的行应采用“密级★保密期限”格式。"""
    for idx, ln in enumerate(lines[:6]):
        m = re.match(r"^(绝密|机密|秘密)", ln)
        if m and "★" not in ln:
            return [Finding(
                "error", "密级标注",
                f"第{idx + 1}行“{ln[:20]}”疑为密级标注，应采用“{m.group(1)}★保密期限”"
                f"格式（如：秘密★1年），并置于首页版心左上角")]
    return []


def _check_attachments(lines: list[str]) -> list[Finding]:
    """附件说明：名称后应使用全角冒号；多个附件须用阿拉伯数字编号。"""
    out = []
    for idx, ln in enumerate(lines):
        if re.match(r"^附件[ \u3000]+\S", ln):
            out.append(Finding(
                "warn", "附件说明",
                f"第{idx + 1}行：附件名称后应使用全角冒号（附件：1. ×××）"))
            break
    colon_lines = [i for i, ln in enumerate(lines) if re.match(r"^附件\s*[:：]", ln)]
    if len(colon_lines) > 1 and not any(
            re.match(r"^附件\s*[:：]\s*\d", lines[i]) for i in colon_lines):
        out.append(Finding(
            "warn", "附件说明",
            f"第{colon_lines[0] + 1}行：多个附件说明应使用阿拉伯数字编号"
            "（附件：1. ×××）"))
    return out


def inspect_docx(path: str) -> list[Finding]:
    """对生成的/外来的 docx 做 GB/T 9704 常用格式检查。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Mm

    out: list[Finding] = []
    doc = Document(path)

    # 页边距
    sec = doc.sections[0]
    got = (round(sec.top_margin / Mm(1)), round(sec.bottom_margin / Mm(1)),
           round(sec.left_margin / Mm(1)), round(sec.right_margin / Mm(1)))
    if got != _GB["margin"]:
        out.append(Finding("info", "页边距",
                           f"当前 上{got[0]} 下{got[1]} 左{got[2]} 右{got[3]}mm；"
                           f"GB/T 9704 常用 上37 下35 左28 右26mm"))

    def east_asia(run) -> str:
        rpr = run._element.rPr
        if rpr is None:
            return ""
        rf = rpr.find(qn("w:rFonts"))
        return rf.get(qn("w:eastAsia")) if rf is not None else ""

    body_n = h_n = 0
    font_bad = size_bad = ls_bad = hfont_bad = 0
    for par in doc.paragraphs:
        text = par.text.strip()
        if not text:
            continue
        style = (par.style.name or "")
        is_heading = style.lower().startswith("heading") or "标题" in style
        if is_heading:
            h_n += 1
            if par.runs:
                fam = east_asia(par.runs[0])
                if fam and fam not in _GB["h1_font"] | _GB["h2_font"] \
                        | _GB["body_font"]:
                    hfont_bad += 1
            continue
        body_n += 1
        if body_n > 200:
            continue
        if par.runs:
            fam = east_asia(par.runs[0])
            size = par.runs[0].font.size.pt if par.runs[0].font.size else None
            if fam and fam not in _GB["body_font"]:
                font_bad += 1
            if size and abs(size - _GB["body_size"]) > 0.6:
                size_bad += 1
        pf = par.paragraph_format
        if pf.line_spacing_rule is not None and pf.line_spacing is not None:
            try:
                if abs(pf.line_spacing.pt - _GB["line_spacing"]) > 1.5:
                    ls_bad += 1
            except AttributeError:
                pass

    if body_n and font_bad / min(body_n, 200) > 0.3:
        out.append(Finding("warn", "正文字体",
                           f"多数正文段落非仿宋类字体（{font_bad} 段）"))
    if body_n and size_bad / min(body_n, 200) > 0.3:
        out.append(Finding("warn", "正文字号",
                           f"多数正文段落非三号（16pt），共 {size_bad} 段"))
    if body_n and ls_bad / min(body_n, 200) > 0.3:
        out.append(Finding("info", "行距",
                           f"多数正文段落非固定 28 磅行距，共 {ls_bad} 段"))
    if h_n and hfont_bad / h_n > 0.5:
        out.append(Finding("info", "标题字体",
                           "多数标题未使用黑体/楷体/仿宋体系，请核对"))

    # 页码域
    footer = doc.sections[0].footer
    has_page = any("PAGE" in p._p.xml for p in footer.paragraphs)
    if not has_page:
        out.append(Finding("info", "页码", "页脚未检测到页码域（PAGE）"))
    return out
